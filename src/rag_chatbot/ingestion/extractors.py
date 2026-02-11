"""Document ingestion from various sources."""

from abc import ABC, abstractmethod
from typing import List, Dict, Any, Optional, BinaryIO
from dataclasses import dataclass
from pathlib import Path
import io
import hashlib

from rag_chatbot.core.logging_config import get_logger
from rag_chatbot.core.exceptions import ChatbotException

logger = get_logger(__name__)


class IngestionError(ChatbotException):
    """Raised when document ingestion fails."""

    pass


@dataclass
class Document:
    """A document with content and metadata."""

    content: str
    metadata: Dict[str, Any]
    source: str
    doc_id: str


class BaseExtractor(ABC):
    """Base class for document extractors."""

    supported_extensions: List[str] = []

    @abstractmethod
    def extract(self, file_obj: BinaryIO, filename: str) -> Document:
        """Extract text from file.

        Args:
            file_obj: File object to read
            filename: Original filename

        Returns:
            Extracted document
        """
        pass

    def _generate_id(self, content: str, filename: str) -> str:
        """Generate unique document ID."""
        hash_input = f"{filename}:{content[:1000]}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]


class PDFExtractor(BaseExtractor):
    """Extract text from PDF files."""

    supported_extensions = [".pdf"]

    def extract(self, file_obj: BinaryIO, filename: str) -> Document:
        """Extract text from PDF."""
        try:
            import PyPDF2

            reader = PyPDF2.PdfReader(file_obj)
            text_parts = []

            for i, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        text_parts.append(f"--- Page {i + 1} ---\n{text}")
                except Exception as e:
                    logger.warning(f"Failed to extract page {i + 1}: {e}")

            content = "\n\n".join(text_parts)

            metadata = {
                "filename": filename,
                "pages": len(reader.pages),
                "type": "pdf",
            }

            return Document(
                content=content,
                metadata=metadata,
                source=filename,
                doc_id=self._generate_id(content, filename),
            )

        except ImportError:
            raise IngestionError("PyPDF2 not installed. Run: pip install PyPDF2")
        except Exception as e:
            raise IngestionError(f"Failed to extract PDF: {e}") from e


class WordExtractor(BaseExtractor):
    """Extract text from Word documents."""

    supported_extensions = [".docx", ".doc"]

    def extract(self, file_obj: BinaryIO, filename: str) -> Document:
        """Extract text from Word document."""
        try:
            import docx

            doc = docx.Document(file_obj)
            text_parts = []

            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    text_parts.append(" | ".join(row_text))

            content = "\n".join(text_parts)

            metadata = {
                "filename": filename,
                "paragraphs": len(doc.paragraphs),
                "type": "word",
            }

            return Document(
                content=content,
                metadata=metadata,
                source=filename,
                doc_id=self._generate_id(content, filename),
            )

        except ImportError:
            raise IngestionError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            raise IngestionError(f"Failed to extract Word doc: {e}") from e


class TextExtractor(BaseExtractor):
    """Extract text from plain text files."""

    supported_extensions = [".txt", ".md", ".rst", ".json", ".xml"]

    def extract(self, file_obj: BinaryIO, filename: str) -> Document:
        """Extract text from plain text file."""
        try:
            content = file_obj.read().decode("utf-8")

            file_type = Path(filename).suffix.lower()

            metadata = {
                "filename": filename,
                "type": file_type.replace(".", ""),
                "characters": len(content),
            }

            return Document(
                content=content,
                metadata=metadata,
                source=filename,
                doc_id=self._generate_id(content, filename),
            )

        except UnicodeDecodeError:
            raise IngestionError(f"File {filename} is not valid UTF-8 text")
        except Exception as e:
            raise IngestionError(f"Failed to extract text: {e}") from e


class CSVExtractor(BaseExtractor):
    """Extract data from CSV/Excel files."""

    supported_extensions = [".csv", ".xlsx", ".xls"]

    def extract(self, file_obj: BinaryIO, filename: str) -> Document:
        """Extract text from CSV/Excel."""
        ext = Path(filename).suffix.lower()

        try:
            if ext == ".csv":
                import csv

                content = file_obj.read().decode("utf-8")
                reader = csv.DictReader(io.StringIO(content))
                rows = list(reader)

                # Convert to readable text
                text_parts = []
                for i, row in enumerate(rows[:1000]):  # Limit rows
                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items())
                    text_parts.append(f"Row {i + 1}: {row_text}")

                content = "\n".join(text_parts)
                metadata = {"filename": filename, "rows": len(rows), "type": "csv"}

            else:  # Excel
                import pandas as pd

                df = pd.read_excel(file_obj)

                # Convert to readable text
                text_parts = []
                for idx, row in df.head(1000).iterrows():
                    row_text = " | ".join(f"{k}: {v}" for k, v in row.items())
                    text_parts.append(f"Row {idx + 1}: {row_text}")

                content = "\n".join(text_parts)
                metadata = {
                    "filename": filename,
                    "rows": len(df),
                    "columns": len(df.columns),
                    "type": "excel",
                }

            return Document(
                content=content,
                metadata=metadata,
                source=filename,
                doc_id=self._generate_id(content, filename),
            )

        except ImportError as e:
            if "pandas" in str(e):
                raise IngestionError("pandas not installed. Run: pip install pandas openpyxl")
            raise
        except Exception as e:
            raise IngestionError(f"Failed to extract spreadsheet: {e}") from e


class WebExtractor(BaseExtractor):
    """Extract text from web pages."""

    supported_extensions = [".url", ".html", ".htm"]

    async def extract_from_url(self, url: str) -> Document:
        """Extract text from URL."""
        try:
            import httpx
            from bs4 import BeautifulSoup

            async with httpx.AsyncClient() as client:
                response = await client.get(url, timeout=30)
                response.raise_for_status()

                soup = BeautifulSoup(response.text, "html.parser")

                # Remove script and style elements
                for script in soup(["script", "style"]):
                    script.decompose()

                # Get text
                text = soup.get_text()

                # Clean up whitespace
                lines = (line.strip() for line in text.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                content = "\n".join(chunk for chunk in chunks if chunk)

                # Get title
                title = soup.find("title")
                title_text = title.get_text() if title else url

                metadata = {
                    "url": url,
                    "title": title_text,
                    "type": "web",
                }

                return Document(
                    content=content,
                    metadata=metadata,
                    source=url,
                    doc_id=self._generate_id(content, url),
                )

        except ImportError:
            raise IngestionError(
                "beautifulsoup4 not installed. Run: pip install beautifulsoup4 httpx"
            )
        except Exception as e:
            raise IngestionError(f"Failed to extract web page: {e}") from e


class DocumentProcessor:
    """Process documents from various sources."""

    def __init__(self):
        """Initialize document processor."""
        self.extractors = {
            ext: extractor()
            for extractor in [
                PDFExtractor,
                WordExtractor,
                TextExtractor,
                CSVExtractor,
            ]
            for ext in extractor.supported_extensions
        }
        self.web_extractor = WebExtractor()

    def process_file(self, file_obj: BinaryIO, filename: str) -> Document:
        """Process a file.

        Args:
            file_obj: File object
            filename: Original filename

        Returns:
            Extracted document
        """
        ext = Path(filename).suffix.lower()

        if ext not in self.extractors:
            raise IngestionError(f"Unsupported file type: {ext}")

        extractor = self.extractors[ext]
        return extractor.extract(file_obj, filename)

    async def process_url(self, url: str) -> Document:
        """Process a URL.

        Args:
            url: Web page URL

        Returns:
            Extracted document
        """
        return await self.web_extractor.extract_from_url(url)

    def process_text(self, text: str, source: str = "manual") -> Document:
        """Process raw text.

        Args:
            text: Raw text content
            source: Source identifier

        Returns:
            Document
        """
        return Document(
            content=text,
            metadata={"type": "text", "source": source},
            source=source,
            doc_id=self._generate_id(text, source),
        )
